"""
AI-First Document Extraction using Azure Document Intelligence
Handles complex document structures that basic parsers miss:
- Tables in DOCX/PDF
- Headers/Footers
- Text in shapes/textboxes
- Multi-column layouts
- Scanned images (OCR)
"""
import logging
from typing import BinaryIO, Dict, Optional
from pathlib import Path
import io

# Basic parsers as fallback
from docx import Document as DocxDocument
import PyPDF2

try:
    from azure.ai.formrecognizer import DocumentAnalysisClient
    from azure.identity import DefaultAzureCredential
    from azure.core.credentials import AzureKeyCredential
    AZURE_DOC_INTEL_AVAILABLE = True
except ImportError:
    AZURE_DOC_INTEL_AVAILABLE = False

try:
    from openpyxl import load_workbook
    EXCEL_SUPPORT = True
except ImportError:
    EXCEL_SUPPORT = False

logger = logging.getLogger(__name__)


class AIDocumentExtractor:
    """AI-First document extraction with fallback to basic parsers"""
    
    def __init__(
        self,
        document_intelligence_endpoint: Optional[str] = None,
        document_intelligence_key: Optional[str] = None,
        use_azure_ad: bool = True
    ):
        """
        Initialize AI Document Extractor
        
        Args:
            document_intelligence_endpoint: Azure Document Intelligence endpoint
            document_intelligence_key: API key (if not using Managed Identity)
            use_azure_ad: Use Azure AD (Managed Identity) authentication
        """
        self.client = None
        
        if AZURE_DOC_INTEL_AVAILABLE and document_intelligence_endpoint:
            try:
                if use_azure_ad:
                    credential = DefaultAzureCredential()
                    self.client = DocumentAnalysisClient(
                        endpoint=document_intelligence_endpoint,
                        credential=credential
                    )
                    logger.info("✅ Azure Document Intelligence initialized with Managed Identity")
                elif document_intelligence_key:
                    credential = AzureKeyCredential(document_intelligence_key)
                    self.client = DocumentAnalysisClient(
                        endpoint=document_intelligence_endpoint,
                        credential=credential
                    )
                    logger.info("✅ Azure Document Intelligence initialized with API key")
            except Exception as e:
                logger.warning(f"⚠️ Failed to initialize Document Intelligence: {e}. Falling back to basic parsers.")
    
    def extract_with_metadata(self, file_content: BinaryIO, filename: str) -> Dict:
        """
        Extract document content with AI-first approach
        
        Args:
            file_content: File binary content
            filename: Original filename
            
        Returns:
            Dict with 'text' and 'pages' metadata
        """
        ext = Path(filename).suffix.lower()
        
        # Try AI extraction first for supported formats
        if self.client and ext in ['.pdf', '.docx', '.doc', '.png', '.jpg', '.jpeg', '.tiff']:
            try:
                return self._extract_with_document_intelligence(file_content, filename)
            except Exception as e:
                logger.warning(f"⚠️ AI extraction failed for {filename}: {e}. Falling back to basic parser.")
        
        # Fallback to basic parsers
        return self._extract_with_basic_parser(file_content, filename)
    
    def _extract_with_document_intelligence(self, file_content: BinaryIO, filename: str) -> Dict:
        """Extract using Azure Document Intelligence (AI-First)"""
        import time
        logger.info(f"🤖 Using AI extraction for: {filename}")
        
        # Reset file pointer
        file_content.seek(0)
        file_bytes = file_content.read()
        file_size_mb = len(file_bytes) / (1024 * 1024)
        
        logger.info(f"📄 Document size: {file_size_mb:.2f} MB")
        
        # Analyze document with "prebuilt-document" model (handles all formats)
        start_time = time.time()
        try:
            poller = self.client.begin_analyze_document(
                "prebuilt-document",
                document=file_bytes
            )
            
            # PRODUCTION FIX: Add timeout based on file size
            # Base: 60s + 30s per MB (reasonable for complex documents)
            timeout_seconds = max(60, int(60 + (file_size_mb * 30)))
            logger.info(f"⏱️  DI timeout set to {timeout_seconds}s for {file_size_mb:.2f}MB file")
            
            result = poller.result(timeout=timeout_seconds)
            elapsed = time.time() - start_time
            logger.info(f"✅ DI analysis completed in {elapsed:.1f}s")
            
        except Exception as e:
            elapsed = time.time() - start_time
            logger.error(f"❌ DI analysis failed after {elapsed:.1f}s: {type(e).__name__}: {str(e)}")
            raise
        
        # Extract text with page information
        pages_metadata = []
        full_text_parts = []
        
        for page_num, page in enumerate(result.pages, start=1):
            page_text_parts = []
            
            # Extract text from lines (maintains reading order)
            if hasattr(page, 'lines'):
                for line in page.lines:
                    page_text_parts.append(line.content)
            
            page_text = "\n".join(page_text_parts)
            
            pages_metadata.append({
                'page_num': page_num,
                'text': page_text,
                'start_char': len("\n\n".join(full_text_parts)),
                'end_char': len("\n\n".join(full_text_parts)) + len(page_text)
            })
            
            full_text_parts.append(page_text)
        
        # Extract tables separately (critical for complex documents)
        table_text_parts = []
        if hasattr(result, 'tables') and result.tables:
            logger.info(f"📊 Found {len(result.tables)} tables in {filename}")
            for table_num, table in enumerate(result.tables, start=1):
                table_text_parts.append(f"\n--- Table {table_num} ---")
                
                # Build table structure
                rows = {}
                for cell in table.cells:
                    row_idx = cell.row_index
                    if row_idx not in rows:
                        rows[row_idx] = {}
                    rows[row_idx][cell.column_index] = cell.content
                
                # Format as text
                for row_idx in sorted(rows.keys()):
                    row_cells = rows[row_idx]
                    row_text = " | ".join([row_cells.get(col, "") for col in sorted(row_cells.keys())])
                    table_text_parts.append(row_text)
        
        # Combine content: pages + tables
        full_text = "\n\n".join(full_text_parts)
        if table_text_parts:
            full_text += "\n\n" + "\n".join(table_text_parts)
        
        logger.info(f"✅ AI extracted {len(full_text)} characters from {filename}")
        
        return {
            'text': full_text,
            'pages': pages_metadata,
            'extraction_method': 'azure_document_intelligence'
        }
    
    def _extract_with_basic_parser(self, file_content: BinaryIO, filename: str) -> Dict:
        """Fallback to basic parsers (python-docx, PyPDF2, etc.)"""
        logger.info(f"📄 Using basic parser for: {filename}")
        
        ext = Path(filename).suffix.lower()
        file_content.seek(0)
        
        if ext == '.pdf':
            return self._extract_pdf_basic(file_content)
        elif ext in ['.docx', '.doc']:
            return self._extract_docx_enhanced(file_content)
        elif ext == '.txt':
            text = file_content.read().decode('utf-8', errors='ignore')
            return {
                'text': text,
                'pages': [{'page_num': 1, 'text': text, 'start_char': 0, 'end_char': len(text)}],
                'extraction_method': 'basic_txt'
            }
        elif ext in ['.xlsx', '.xls']:
            return self._extract_excel_basic(file_content)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _extract_docx_enhanced(self, file_content: BinaryIO) -> Dict:
        """Enhanced DOCX extraction - includes tables!"""
        try:
            doc = DocxDocument(file_content)
            text_parts = []
            pages_metadata = []
            current_char = 0
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables (THIS WAS MISSING!)
            logger.info(f"📊 Extracting {len(doc.tables)} tables from DOCX")
            for table_num, table in enumerate(doc.tables, start=1):
                text_parts.append(f"\n--- Table {table_num} ---")
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            
            # Create page metadata (treat each section as a page)
            pages_metadata.append({
                'page_num': 1,
                'text': full_text,
                'start_char': 0,
                'end_char': len(full_text)
            })
            
            logger.info(f"✅ Enhanced DOCX extraction: {len(full_text)} characters")
            
            return {
                'text': full_text,
                'pages': pages_metadata,
                'extraction_method': 'enhanced_docx'
            }
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX: {str(e)}")
    
    def _extract_pdf_basic(self, file_content: BinaryIO) -> Dict:
        """Basic PDF extraction"""
        try:
            pdf_reader = PyPDF2.PdfReader(file_content)
            pages_metadata = []
            full_text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                page_text = page.extract_text()
                if page_text.strip():
                    pages_metadata.append({
                        'page_num': page_num,
                        'text': page_text,
                        'start_char': len("\n\n".join(full_text_parts)),
                        'end_char': len("\n\n".join(full_text_parts)) + len(page_text)
                    })
                    full_text_parts.append(page_text)
            
            full_text = "\n\n".join(full_text_parts)
            
            return {
                'text': full_text,
                'pages': pages_metadata,
                'extraction_method': 'basic_pdf'
            }
        except Exception as e:
            raise ValueError(f"Failed to extract PDF: {str(e)}")
    
    def _extract_excel_basic(self, file_content: BinaryIO) -> Dict:
        """Basic Excel extraction"""
        if not EXCEL_SUPPORT:
            raise ValueError("Excel support not available. Install openpyxl.")
        
        try:
            workbook = load_workbook(file_content, data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"\n--- Sheet: {sheet_name} ---")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            
            return {
                'text': full_text,
                'pages': [{'page_num': 1, 'text': full_text, 'start_char': 0, 'end_char': len(full_text)}],
                'extraction_method': 'basic_excel'
            }
        except Exception as e:
            raise ValueError(f"Failed to extract Excel: {str(e)}")
