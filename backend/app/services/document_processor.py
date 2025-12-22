"""Document processing service for extracting text from various document types"""
import io
import os
from pathlib import Path
from typing import BinaryIO, Optional
import PyPDF2
from docx import Document as DocxDocument
from app.services.ai_document_extractor import AIDocumentExtractor
from app.config import settings
import logging

logger = logging.getLogger(__name__)

# Optional imports with fallback
try:
    from openpyxl import load_workbook
    EXCEL_SUPPORT = True
except ImportError:
    EXCEL_SUPPORT = False

try:
    from PIL import Image
    import pytesseract
    IMAGE_SUPPORT = True
except ImportError:
    IMAGE_SUPPORT = False


class DocumentProcessor:
    """Handles document parsing and text extraction with AI-First approach"""
    
    SUPPORTED_TYPES = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".txt": "text/plain",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
    }
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document processor with AI-First extraction
        
        Args:
            chunk_size: Target size for text chunks (in characters)
            chunk_overlap: Overlap between chunks (in characters)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize AI Document Extractor if configured
        self.ai_extractor = None
        if settings.use_document_intelligence:
            try:
                self.ai_extractor = AIDocumentExtractor(
                    document_intelligence_endpoint=settings.azure_document_intelligence_endpoint,
                    document_intelligence_key=settings.azure_document_intelligence_key,
                    use_azure_ad=settings.use_azure_ad_auth
                )
                logger.info("✅ AI Document Extractor initialized")
            except Exception as e:
                logger.warning(f"⚠️ Failed to initialize AI extractor: {e}. Using basic parsers.")
    
    def is_supported(self, filename: str) -> bool:
        """Check if file type is supported"""
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_TYPES
    
    def get_file_type(self, filename: str) -> str:
        """Get MIME type for file"""
        ext = Path(filename).suffix.lower()
        return self.SUPPORTED_TYPES.get(ext, "application/octet-stream")
    
    def extract_text(self, file_content: BinaryIO, filename: str) -> str:
        """
        Extract text from document (backward compatible)
        
        Args:
            file_content: File binary content
            filename: Original filename to determine type
            
        Returns:
            Extracted text content
        """
        structured_data = self.extract_with_metadata(file_content, filename)
        return structured_data['text']
    
    def extract_with_metadata(self, file_content: BinaryIO, filename: str) -> dict:
        """
        Extract text with metadata using AI-First approach
        
        Args:
            file_content: File binary content
            filename: Original filename to determine type
            
        Returns:
            Dict with 'text', 'pages' metadata, and 'extraction_method'
        """
        # Try AI extraction first if available
        if self.ai_extractor:
            try:
                logger.info(f"🤖 Attempting AI extraction for: {filename}")
                result = self.ai_extractor.extract_with_metadata(file_content, filename)
                logger.info(f"✅ AI extraction successful: {result.get('extraction_method')}")
                return result
            except Exception as e:
                logger.warning(f"⚠️ AI extraction failed: {e}. Falling back to basic parser.")
                file_content.seek(0)  # Reset file pointer
        
        # Fallback to basic parsers
        logger.info(f"📄 Using basic parser for: {filename}")
        ext = Path(filename).suffix.lower()
        
        if ext == ".pdf":
            return self._extract_pdf_with_pages(file_content)
        elif ext in [".docx", ".doc"]:
            return self._extract_docx_with_paragraphs(file_content)
        elif ext == ".txt":
            text = self._extract_txt(file_content)
            return {
                'text': text,
                'pages': [{'page_num': 1, 'text': text, 'start_char': 0, 'end_char': len(text)}]
            }
        elif ext in [".xlsx", ".xls"]:
            return self._extract_excel(file_content, ext)
        elif ext in [".png", ".jpg", ".jpeg"]:
            return self._extract_image(file_content)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _extract_pdf(self, file_content: BinaryIO) -> str:
        """Extract text from PDF file (backward compatible)"""
        data = self._extract_pdf_with_pages(file_content)
        return data['text']
    
    def _extract_pdf_with_pages(self, file_content: BinaryIO) -> dict:
        """Extract text from PDF with page metadata"""
        try:
            pdf_reader = PyPDF2.PdfReader(file_content)
            text_parts = []
            pages_metadata = []
            current_char = 0
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        page_text = f"[Page {page_num + 1}]\n{text}"
                        text_parts.append(page_text)
                        
                        # Track page metadata
                        pages_metadata.append({
                            'page_num': page_num + 1,
                            'text': text,
                            'start_char': current_char,
                            'end_char': current_char + len(page_text)
                        })
                        current_char += len(page_text) + 2  # +2 for "\n\n"
                    
                    # Force garbage collection every 10 pages to prevent memory buildup
                    if (page_num + 1) % 10 == 0:
                        import gc
                        gc.collect()
                        
                except Exception as e:
                    print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
                    continue
            
            full_text = "\n\n".join(text_parts)
            result = {
                'text': full_text,
                'pages': pages_metadata
            }
            
            # Final cleanup
            del pdf_reader, text_parts, pages_metadata
            import gc
            gc.collect()
            
            return result
        except Exception as e:
            raise ValueError(f"Failed to extract PDF: {str(e)}")
    
    def _extract_docx(self, file_content: BinaryIO) -> str:
        """Extract text from DOCX file (backward compatible)"""
        data = self._extract_docx_with_paragraphs(file_content)
        return data['text']
    
    def _extract_docx_with_paragraphs(self, file_content: BinaryIO) -> dict:
        """Extract text from DOCX with paragraph metadata - NOW INCLUDES TABLES!"""
        try:
            doc = DocxDocument(file_content)
            text_parts = []
            pages_metadata = []
            current_char = 0
            
            # Extract paragraphs
            for para_num, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    para_text = para.text
                    text_parts.append(para_text)
                    
                    # Track paragraph as "page" for consistency
                    pages_metadata.append({
                        'page_num': para_num + 1,  # Paragraph number
                        'text': para_text,
                        'start_char': current_char,
                        'end_char': current_char + len(para_text)
                    })
                    current_char += len(para_text) + 2  # +2 for "\n\n"
            
            # CRITICAL FIX: Extract tables (was completely missing!)
            if doc.tables:
                logger.info(f"📊 Extracting {len(doc.tables)} tables from DOCX")
                text_parts.append("\n=== TABLES ===")
                
                for table_num, table in enumerate(doc.tables, start=1):
                    text_parts.append(f"\n--- Table {table_num} ---")
                    
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            text_parts.append(row_text)
                    
                    text_parts.append("")  # Blank line between tables
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"✅ Enhanced DOCX extraction: {len(full_text)} chars, {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
            
            return {
                'text': full_text,
                'pages': pages_metadata,
                'extraction_method': 'enhanced_docx_with_tables'
            }
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX: {str(e)}")
    
    def _extract_txt(self, file_content: BinaryIO) -> str:
        """Extract text from TXT file"""
        try:
            content = file_content.read()
            # Try UTF-8 first, fall back to latin-1
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                return content.decode('latin-1')
        except Exception as e:
            raise ValueError(f"Failed to extract TXT: {str(e)}")
    
    def chunk_text(self, text: str, metadata: Optional[dict] = None, pages_info: Optional[list] = None) -> list[dict]:
        """
        Split text into overlapping chunks with page number tracking
        
        Args:
            text: Full text to chunk
            metadata: Optional metadata to attach to each chunk
            pages_info: Optional list of page metadata from extract_with_metadata
            
        Returns:
            List of chunks with metadata including page numbers
        """
        if not text.strip():
            return []
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            # Calculate end position
            end = start + self.chunk_size
            
            # If not at the end, try to break at sentence or word boundary
            if end < len(text):
                # Look for sentence end (., !, ?)
                for delimiter in ['. ', '! ', '? ', '\n\n', '\n']:
                    last_delim = text.rfind(delimiter, start, end)
                    if last_delim != -1:
                        end = last_delim + len(delimiter)
                        break
            
            # Extract chunk
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunk_data = {
                    "text": chunk_text,
                    "chunk_index": chunk_index,
                    "start_char": start,
                    "end_char": end
                }
                
                # Determine page number(s) for this chunk
                if pages_info:
                    chunk_pages = []
                    for page in pages_info:
                        # Check if chunk overlaps with this page
                        if not (end <= page['start_char'] or start >= page['end_char']):
                            chunk_pages.append(page['page_num'])
                    
                    if chunk_pages:
                        chunk_data['page_number'] = chunk_pages[0]  # Primary page
                        chunk_data['page_numbers'] = chunk_pages  # All pages it spans
                
                if metadata:
                    chunk_data.update(metadata)
                
                chunks.append(chunk_data)
                chunk_index += 1
            
            # Move start position (with overlap)
            start = end - self.chunk_overlap if end < len(text) else end
        
        return chunks
    
    def _extract_excel(self, file_content: BinaryIO, ext: str) -> dict:
        """Extract text from Excel files"""
        if not EXCEL_SUPPORT:
            raise ValueError("Excel support not available. Install openpyxl and xlrd.")
        
        try:
            workbook = load_workbook(file_content, read_only=True, data_only=True)
            sheets_text = []
            sheets_metadata = []
            current_char = 0
            
            for sheet_idx, sheet_name in enumerate(workbook.sheetnames):
                sheet = workbook[sheet_name]
                rows_text = []
                
                # Extract all cell values
                for row in sheet.iter_rows(values_only=True):
                    # Filter out None values and convert to strings
                    row_values = [str(cell) for cell in row if cell is not None]
                    if row_values:
                        rows_text.append(" | ".join(row_values))
                
                if rows_text:
                    sheet_text = f"[Sheet: {sheet_name}]\n" + "\n".join(rows_text)
                    sheets_text.append(sheet_text)
                    
                    sheets_metadata.append({
                        'page_num': sheet_idx + 1,
                        'text': sheet_text,
                        'start_char': current_char,
                        'end_char': current_char + len(sheet_text),
                        'sheet_name': sheet_name
                    })
                    current_char += len(sheet_text) + 2
            
            full_text = "\n\n".join(sheets_text)
            return {
                'text': full_text,
                'pages': sheets_metadata
            }
        except Exception as e:
            raise ValueError(f"Failed to extract Excel: {str(e)}")
    
    def _extract_image(self, file_content: BinaryIO) -> dict:
        """Extract text from images using OCR with memory optimization"""
        if not IMAGE_SUPPORT:
            raise ValueError("Image/OCR support not available. Install Pillow and pytesseract.")
        
        image = None
        try:
            image = Image.open(file_content)
            
            # Resize large images to prevent OOM (max 4000px on longest side)
            max_dimension = 4000
            if max(image.size) > max_dimension:
                ratio = max_dimension / max(image.size)
                new_size = tuple(int(dim * ratio) for dim in image.size)
                image = image.resize(new_size, Image.Resampling.LANCZOS)
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            if not text.strip():
                text = "[No text detected in image]"
            
            result = {
                'text': text,
                'pages': [{'page_num': 1, 'text': text, 'start_char': 0, 'end_char': len(text)}]
            }
            
            return result
        except Exception as e:
            raise ValueError(f"Failed to extract text from image: {str(e)}")
        finally:
            # Explicitly close and delete image to free memory
            if image is not None:
                image.close()
                del image
            import gc
            gc.collect()
